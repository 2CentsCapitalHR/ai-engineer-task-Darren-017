from pathlib import Path
import re
from typing import List, Dict
from docx import Document

def annotate_docx(src: Path, issues: List[Dict], out_path: Path):
    doc = Document(str(src))

    counter = 1
    for issue in issues:
        anchor = issue.get("anchor_regex")
        if not anchor:
            continue
        pattern = re.compile(anchor, re.I)
        placed = False
        for para in doc.paragraphs:
            if pattern.search(para.text or ""):
                note = f" [COMMENT {counter}: {issue.get('issue')} | Suggestion: {issue.get('suggestion')}]"
                para.text = (para.text or "") + note
                counter += 1
                placed = True
                break
        if not placed:
            continue

    if issues:
        doc.add_paragraph("")
        doc.add_paragraph("=== REVIEW NOTES (Auto-generated) ===")
        for i, issue in enumerate(issues, start=1):
            line = f"[{i}] {issue.get('section')}: {issue.get('issue')} (Severity: {issue.get('severity')})."
            if issue.get("suggestion"):
                line += f" Suggestion: {issue['suggestion']}"
            doc.add_paragraph(line)

    doc.save(str(out_path))
